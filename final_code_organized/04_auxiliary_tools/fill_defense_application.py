from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"D:/field_batch_output_compressed_air")
SOURCE = ROOT / "defense_application_original.docx"
OUTPUT = ROOT / "深港微电子学院本科毕业论文答辩申请表_学生填写完成.docx"


SELF_STATEMENT = (
    "本人已按本科毕业论文工作计划完成题为《深度学习驱动超表面设计及其在红外探测中的应用》的毕业论文。"
    "论文围绕红外 MIM 超表面吸收器设计中全波仿真开销大、结构搜索效率低等问题，完成了器件结构参数化建模、"
    "仿真数据集构建、基于深度学习的前向代理模型训练，以及结合进化算法的目标导向逆向设计与结果分析。"
    "论文撰写过程中已根据指导老师和评阅意见，对研究内容表述、结果讨论、格式规范及参考文献等进行了修改完善，"
    "相关材料经指导老师审核检查，达到本科毕业论文答辩要求。现申请参加毕业论文答辩，请予批准。"
)


def clear_cell(cell):
    for paragraph in list(cell.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False):
    paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.3
    if first_line:
        pf.first_line_indent = Pt(21)
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def add_text(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False):
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(text)
    style_paragraph(paragraph, align=align, first_line=first_line)
    return paragraph


def main():
    doc = Document(SOURCE)
    table = doc.tables[0]

    statement_cell = table.rows[3].cells[1]
    clear_cell(statement_cell)
    set_cell_margins(statement_cell)

    add_text(statement_cell, SELF_STATEMENT, first_line=True)
    add_text(statement_cell, "")
    add_text(statement_cell, "申请人：吕晗熙", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(statement_cell, "日期：2026年5月25日", align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
